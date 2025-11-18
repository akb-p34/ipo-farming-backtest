"""
Script to convert the Backtest Plan from Markdown to Word (.docx) and PDF formats
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def parse_markdown_to_word(md_file_path, docx_file_path):
    """
    Parse markdown file and create a formatted Word document
    """
    # Read markdown content
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create Word document
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Process content line by line
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines at start, but preserve spacing elsewhere
        if not line:
            if len(doc.paragraphs) > 0:
                doc.add_paragraph()
            i += 1
            continue

        # Heading 1 (# )
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Heading 2 (## )
        elif line.startswith('## ') and not line.startswith('### '):
            text = line[3:].strip()
            doc.add_heading(text, level=2)

        # Heading 3 (### )
        elif line.startswith('### '):
            text = line[4:].strip()
            doc.add_heading(text, level=3)

        # Horizontal rule (---)
        elif line.strip() == '---':
            p = doc.add_paragraph()
            p.add_run('_' * 80)

        # Bold text (**text**)
        elif '**' in line:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        # Bullet points (- )
        elif line.startswith('- '):
            text = line[2:].strip()
            # Handle bold within bullets
            p = doc.add_paragraph(style='List Bullet')
            if '**' in text:
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)
            else:
                p.add_run(text)

        # Numbered lists (1. )
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line).strip()
            p = doc.add_paragraph(style='List Number')
            if '**' in text:
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)
            else:
                p.add_run(text)

        # Regular paragraph
        else:
            # Handle bold within regular text
            if '**' in line:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)
            else:
                doc.add_paragraph(line)

        i += 1

    # Save document
    doc.save(docx_file_path)
    print(f"Word document created: {docx_file_path}")

def convert_word_to_pdf(docx_file_path, pdf_file_path):
    """
    Convert Word document to PDF
    """
    try:
        # Try using docx2pdf library
        from docx2pdf import convert
        convert(docx_file_path, pdf_file_path)
        print(f"PDF created: {pdf_file_path}")
        return True
    except ImportError:
        print("docx2pdf not installed. Trying alternative methods...")

        # Try using system commands
        import subprocess
        import platform

        system = platform.system()

        if system == 'Darwin':  # macOS
            try:
                # Try using textutil (built-in on macOS)
                subprocess.run([
                    'textutil',
                    '-convert', 'pdf',
                    docx_file_path,
                    '-output', pdf_file_path
                ], check=True)
                print(f"PDF created using textutil: {pdf_file_path}")
                return True
            except subprocess.CalledProcessError:
                print("textutil failed.")

            try:
                # Try using soffice (LibreOffice) if installed
                subprocess.run([
                    'soffice',
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', os.path.dirname(pdf_file_path),
                    docx_file_path
                ], check=True)
                print(f"PDF created using LibreOffice: {pdf_file_path}")
                return True
            except (subprocess.CalledProcessError, FileNotFoundError):
                print("LibreOffice not found.")

        print("\nCould not create PDF automatically.")
        print("Please open the Word document and use 'Save As PDF' manually.")
        print(f"Word document location: {docx_file_path}")
        return False

if __name__ == "__main__":
    import os

    # File paths
    base_dir = "/Users/akbarpathan/Desktop/Dev/Trading Expirements/IPO Farming"
    md_file = os.path.join(base_dir, "IPO_Backtest_Plan.md")
    docx_file = os.path.join(base_dir, "IPO_Backtest_Plan.docx")
    pdf_file = os.path.join(base_dir, "IPO_Backtest_Plan.pdf")

    # Convert markdown to Word
    print("Creating Word document...")
    parse_markdown_to_word(md_file, docx_file)

    # Convert Word to PDF
    print("\nCreating PDF...")
    success = convert_word_to_pdf(docx_file, pdf_file)

    if not success:
        print("\n" + "="*60)
        print("MANUAL PDF CONVERSION INSTRUCTIONS:")
        print("="*60)
        print(f"1. Open: {docx_file}")
        print("2. Click File > Export > Export as PDF")
        print(f"3. Save as: {pdf_file}")
        print("="*60)
